#!/usr/bin/env python3
"""Generate the connected project-charter Word sample from structured JSON."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SAMPLE_DIR = ROOT / "samples" / "car-control-app"
OUTPUT = SAMPLE_DIR / "汽车手机端控制应用-项目章程示例.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
MID_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_BLUE = "F3F7FB"
INK = "243447"
MUTED = "5B6875"
RED = "B42318"
PALE_RED = "FDECEC"
AMBER = "9A6700"
PALE_AMBER = "FFF8E1"
GREEN = "257A4A"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
CJK_FONT = "Arial Unicode MS"


def load_json(name: str):
    with (SAMPLE_DIR / name).open("r", encoding="utf-8") as handle:
        return json.load(handle)


charter = load_json("project-charter.json")
tree = load_json("reasoning-tree.json")
interview = load_json("interview.json")
fields = charter["fields"]


def set_run_font(run, size=None, bold=None, color=None, italic=None, east_asia=CJK_FONT):
    run.font.name = CJK_FONT
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_fonts = run._element.get_or_add_rPr().rFonts
    r_fonts.set(qn("w:ascii"), CJK_FONT)
    r_fonts.set(qn("w:hAnsi"), CJK_FONT)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), CJK_FONT)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Pt(width_dxa / 20)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color="CBD5E1", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_table(document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], small=False, tiny=False):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    repeat_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, PALE_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(text))
        set_run_font(run, 8 if tiny else 8.5 if small else 9, True, NAVY)
    for row_values in rows:
        row = table.add_row()
        keep_row_together(row)
        for index, value in enumerate(row_values):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, 7.25 if tiny else 8 if small else 8.5, False, INK)
    set_table_geometry(table, widths_dxa)
    if tiny:
        set_table_geometry(table, widths_dxa, indent_dxa=90)
        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell, top=45, start=90, bottom=45, end=90)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_field(paragraph, field_code: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, value, end])
    set_run_font(run, 8, False, MUTED)


def install_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1
    result = {}
    for kind, fmt, text_value, font in (
        ("bullet", "bullet", "•", "Symbol"),
        ("number", "decimal", "%1.", "Calibri"),
    ):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abstract))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        level.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        level.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "480")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "480")
        ind.set(qn("w:hanging"), "240")
        p_pr.append(ind)
        level.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font)
        r_fonts.set(qn("w:hAnsi"), font)
        r_pr.append(r_fonts)
        level.append(r_pr)
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(next_abstract))
        num.append(abstract_ref)
        numbering.append(num)
        result[kind] = next_num
        next_abstract += 1
        next_num += 1
    return result


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_list(document, values, numbered=False, compact=False):
    num_id = NUMBERING["number" if numbered else "bullet"]
    for value in values:
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, num_id)
        paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(str(value))
        set_run_font(run, 10 if compact else 10.5, False, INK)


def add_checklist(document, values):
    for value in values:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(15)
        paragraph.paragraph_format.first_line_indent = Pt(-10)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.12
        box = paragraph.add_run("□  ")
        set_run_font(box, 10, False, MUTED)
        run = paragraph.add_run(str(value))
        set_run_font(run, 10, False, INK)


def add_label_value(document, label: str, envelope, status_note=True):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    label_run = paragraph.add_run(f"{label}　")
    set_run_font(label_run, 10.5, True, MID_BLUE)
    value = envelope.get("value")
    if value in (None, "", []):
        value = "待确认"
    value_run = paragraph.add_run(str(value))
    set_run_font(value_run, 10.5, False, RED if envelope["status"] == "missing" else INK)
    if status_note and envelope["status"] in {"inferred", "assumed", "missing"}:
        labels = {"inferred": "推断", "assumed": "暂定", "missing": "待确认"}
        note_run = paragraph.add_run(f" 〔{labels[envelope['status']]}〕")
        color = RED if envelope["status"] == "missing" else AMBER if envelope["status"] == "assumed" else BLUE
        set_run_font(note_run, 8.5, False, color, True)
    return paragraph


def add_callout(document, title: str, text: str, fill=LIGHT_BLUE, border=BLUE, title_color=NAVY):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(8)
    paragraph.paragraph_format.right_indent = Pt(8)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    borders.append(left)
    p_pr.append(borders)
    title_run = paragraph.add_run(f"{title}  ")
    set_run_font(title_run, 10.5, True, title_color)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, 10, False, INK)
    return paragraph


def add_heading(document, text: str, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_subheading(document, text: str):
    paragraph = document.add_heading(text, level=2)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = CJK_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, MID_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = CJK_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.space_after = Pt(0)
    brand = header_para.add_run("PM KNOWLEDGE LAB")
    set_run_font(brand, 8, True, BLUE)
    descriptor = header_para.add_run("   /   项目章程样本")
    set_run_font(descriptor, 8, False, MUTED)
    p_pr = header_para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_after = Pt(0)
    foot = footer_para.add_run("汽车手机端控制应用项目｜推演草案  ·  第 ")
    set_run_font(foot, 8, False, MUTED)
    add_field(footer_para, "PAGE")
    foot2 = footer_para.add_run(" 页")
    set_run_font(foot2, 8, False, MUTED)

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    document.core_properties.title = "汽车手机端控制应用项目章程（推演草案）"
    document.core_properties.subject = "项目章程连续样本：访谈、推理与模板自动带入"
    document.core_properties.author = "项目管理知识网页"
    document.core_properties.keywords = "项目章程, 项目管理, 汽车应用, 推理树"


document = Document()
configure_document(document)
NUMBERING = install_numbering(document)

# Title block.
title = document.add_paragraph()
title.paragraph_format.space_before = Pt(30)
title.paragraph_format.space_after = Pt(5)
title_run = title.add_run("项目章程")
set_run_font(title_run, 24, True, NAVY)
subtitle = document.add_paragraph()
subtitle.paragraph_format.space_after = Pt(22)
sub_run = subtitle.add_run("汽车手机端控制应用项目  /  推演草案")
set_run_font(sub_run, 13, False, BLUE)

metadata_rows = [
    ["文档版本", str(fields["charter_version"]["value"]), "文档状态", str(fields["document_status"]["value"])],
    ["编制日期", str(fields["prepared_date"]["value"]), "批准日期", "待确认"],
    ["项目经理", "待确认（角色已识别）", "发起人", "待确认（角色：数字产品中心负责人）"],
]
metadata = add_table(document, ["", "", "", ""], metadata_rows, [1350, 3330, 1350, 3330])
# Remove visual header row because this is a label/value data table, not a repeated record table.
metadata._tbl.remove(metadata.rows[0]._tr)
for row in metadata.rows:
    for index, cell in enumerate(row.cells):
        if index in (0, 2):
            set_cell_shading(cell, PALE_BLUE)
            for run in cell.paragraphs[0].runs:
                set_run_font(run, 8.5, True, NAVY)
        elif "待确认" in cell.text:
            for run in cell.paragraphs[0].runs:
                set_run_font(run, 8.5, False, RED)

add_callout(
    document,
    "草案状态",
    "本文件由样本访谈与规则推演生成，供项目启动讨论使用。批准预算、资金来源、项目经理姓名及发起人姓名尚未确认，因此当前不可作为正式批准章程。",
    fill=PALE_RED,
    border=RED,
    title_color=RED,
)

decision = next(node for node in tree["nodes"] if node["id"] == "D-LIFECYCLE")
add_callout(
    document,
    "建议的开发与治理方式",
    f"{decision['title']}。{decision['rationale']}",
    fill=LIGHT_BLUE,
    border=BLUE,
)

# 1.
add_heading(document, "1. 项目目的与预期价值")
add_label_value(document, "项目目的", fields["project_purpose"])
add_label_value(document, "业务需求", fields["business_need"])
add_subheading(document, "预期价值")
add_list(document, fields["expected_value"]["value"])

# 2.
add_heading(document, "2. 可测量的项目目标与成功标准")
objective_rows = [
    [item["objective"], f"{item['metric']}\n{item['target']}", item["deadline"], item["verificationMethod"], item["ownerRole"]]
    for item in fields["measurable_objectives"]["value"]
]
add_table(document, ["目标", "指标与目标值", "期限", "验证方式", "责任角色"], objective_rows, [1650, 2050, 1100, 2860, 1700], small=True)
add_callout(document, "确认提示", "目标值和期限来自当前样本假设，进入章程批准前应由发起人、产品、安全和技术责任角色共同确认。", fill=PALE_AMBER, border=AMBER, title_color=AMBER)

# 3.
add_heading(document, "3. 高层级范围、需求与交付成果")
add_subheading(document, "高层级需求")
add_list(document, fields["high_level_requirements"]["value"], compact=True)
add_label_value(document, "范围边界", fields["scope_boundary"])
add_subheading(document, "主要交付成果")
deliverable_rows = [[item["name"], item["description"], item["acceptanceOwner"]] for item in fields["key_deliverables"]["value"]]
add_table(document, ["交付成果", "说明", "验收责任角色"], deliverable_rows, [1850, 5210, 2300], small=True)
add_subheading(document, "明确不在第一阶段范围内")
add_list(document, fields["out_of_scope"]["value"], compact=True)

# 4.
add_heading(document, "4. 总体里程碑进度计划")
milestone_rows = [[item["name"], item["targetDate"], item["exitCriteria"], item["ownerRole"]] for item in fields["milestone_plan"]["value"]]
add_table(document, ["里程碑", "目标日期", "退出条件", "责任角色"], milestone_rows, [1900, 1200, 4260, 2000], small=True)

# 5.
add_heading(document, "5. 整体项目风险")
risk_rows = [[item["title"], item["rating"], item["ownerRole"], item["responseSummary"]] for item in fields["overall_risk_summary"]["value"]]
add_table(document, ["风险", "等级", "责任角色", "高层级应对"], risk_rows, [2720, 800, 1840, 4000], small=True)
add_callout(document, "风险治理", "章程仅呈现高层摘要。完整的原因—事件—影响、概率与影响评分、触发条件、应急计划及剩余风险保存在风险登记册中。")

# 6.
add_heading(document, "6. 财务资源与关键干系人")
add_callout(document, "预算待确认", "批准预算总额、资金来源和高层级预算分配均未提供。系统未生成虚构金额；该缺口会阻止章程进入“已批准”状态。", fill=PALE_RED, border=RED, title_color=RED)
stakeholder_rows = []
for item in fields["key_stakeholders"]["value"]:
    status_text = "姓名待确认" if item["status"] == "missing-name" else "角色已识别"
    stakeholder_rows.append([item["role"], item["relationship"], status_text])
add_table(document, ["关键角色", "与项目的关系", "确认状态"], stakeholder_rows, [2500, 4800, 2060], small=True)

# 7.
add_heading(document, "7. 审批要求与退出标准")
add_subheading(document, "验收要求")
add_list(document, fields["acceptance_requirements"]["value"], compact=True)
add_subheading(document, "审批流程")
add_list(document, fields["approval_process"]["value"], numbered=True, compact=True)
add_label_value(document, "最终批准角色", fields["final_approver"])
add_subheading(document, "项目退出标准")
add_checklist(document, fields["exit_criteria"]["value"])

# 8.
add_heading(document, "8. 项目经理与发起人权限")
add_subheading(document, "项目经理")
add_label_value(document, "人员", fields["project_manager"])
add_label_value(document, "批准结论", fields["approval_decision"], status_note=False)
add_subheading(document, "主要职责")
add_list(document, fields["project_manager_responsibilities"]["value"], compact=True)
add_subheading(document, "职权边界")
add_list(document, fields["project_manager_authority"]["value"], compact=True)
add_subheading(document, "项目发起人")
add_label_value(document, "人员", fields["sponsor"])
add_subheading(document, "批准权限")
add_list(document, fields["sponsor_authority"]["value"], compact=True)

# Appendix A.
document.add_page_break()
add_heading(document, "附录 A：自动带入与推理追踪")
intro = document.add_paragraph("本附录用于说明内容是如何形成的。正文保持正式表达；证据状态和推理路径在此集中呈现，便于评审者复核。")
intro.paragraph_format.space_after = Pt(8)
status_labels = {"confirmed": "已确认", "inferred": "推断", "assumed": "暂定", "missing": "待确认"}
trace_rows = []
for trace in charter["autofillTrace"]:
    facts = "、".join(trace.get("inputFacts", [])) or "—"
    reasoning = "、".join(trace.get("reasoningIds", [])) or "直接映射"
    trace_rows.append([trace["targetField"], facts, reasoning, status_labels[trace["status"]]])
add_table(document, ["章程字段", "输入事实", "推理/规则", "状态"], trace_rows, [1850, 2860, 3300, 1350], small=True, tiny=True)

add_subheading(document, "生命周期判断的证据链")
reasoning_rows = []
for node in tree["nodes"]:
    if node["id"] in {"F-SAFETY", "F-PRIVACY", "F-FEEDBACK", "F-INTEGRATION", "J-GOVERNANCE", "J-ITERATION", "J-INTEGRATION", "D-LIFECYCLE"}:
        reasoning_rows.append([node["id"], node["type"], node["title"], "、".join(node.get("sourceIds", []))])
add_table(document, ["节点", "类型", "判断内容", "来源"], reasoning_rows, [1450, 1100, 4700, 2110], small=True, tiny=True)

# Appendix B.
document.add_page_break()
add_heading(document, "附录 B：待确认事项与批准阻断项")
add_callout(document, "当前结论", f"本章程可以导出草案，但不可批准。共识别 {len(charter['validation']['blockingIssues'])} 个批准阻断项。", fill=PALE_RED, border=RED, title_color=RED)
question_rows = [[item["id"], item["question"], item["blocks"]] for item in interview["openQuestions"]]
add_table(document, ["编号", "需要确认的问题", "阻断对象"], question_rows, [1200, 6110, 2050], small=True)
add_subheading(document, "系统校验结果")
add_checklist(document, charter["validation"]["blockingIssues"])
closing = document.add_paragraph()
closing.paragraph_format.space_before = Pt(20)
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run("— 章节正文结束 / 后续确认结果应回填到结构化字段后重新导出 —")
set_run_font(run, 9, False, MUTED, True)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
document.save(OUTPUT)
print(OUTPUT)
